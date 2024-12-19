import tkinter as tk
from tkinter import ttk, filedialog, messagebox, colorchooser
import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx2pdf import convert
import pandas as pd
import re
from pathlib import Path


class KeywordFormatDialog:
    def __init__(self, parent, current_format=None):
        self.dialog = tk.Toplevel(parent)
        self.dialog.title("Edit Keyword Format")
        self.dialog.geometry("400x450")

        # Make dialog modal
        self.dialog.transient(parent)
        self.dialog.grab_set()

        # Format settings
        self.format_settings = current_format or {
            'font_name': 'Arial',
            'font_size': 11,
            'font_color': '#000000',
            'bold': False,
            'italic': False,
            'underline': False
        }

        self.create_widgets()

    def create_widgets(self):
        # Font Type
        font_frame = ttk.LabelFrame(
            self.dialog, text="Font Settings", padding=10)
        font_frame.pack(fill="x", padx=20, pady=10)

        ttk.Label(font_frame, text="Font Type:").grid(
            row=0, column=0, padx=5, pady=5)
        self.font_combo = ttk.Combobox(font_frame, values=[
            'Arial', 'Times New Roman', 'Calibri', 'Cambria', 'Georgia',
            'Helvetica', 'Verdana', 'Tahoma'
        ])
        self.font_combo.set(self.format_settings['font_name'])
        self.font_combo.grid(row=0, column=1, padx=5, pady=5)

        # Font Size
        ttk.Label(font_frame, text="Font Size:").grid(
            row=1, column=0, padx=5, pady=5)
        self.size_var = tk.StringVar(
            value=str(self.format_settings['font_size']))
        size_spinbox = ttk.Spinbox(
            font_frame, from_=6, to=72, textvariable=self.size_var)
        size_spinbox.grid(row=1, column=1, padx=5, pady=5)

        # Font Color
        ttk.Label(font_frame, text="Font Color:").grid(
            row=2, column=0, padx=5, pady=5)
        self.color_button = ttk.Button(
            font_frame,
            text="Choose Color",
            command=self.choose_color
        )
        self.color_button.grid(row=2, column=1, padx=5, pady=5)
        self.update_color_button()

        # Style Options Frame
        style_frame = ttk.LabelFrame(
            self.dialog, text="Style Options", padding=10)
        style_frame.pack(fill="x", padx=20, pady=10)

        # Checkboxes for bold, italic, underline
        self.bold_var = tk.BooleanVar(value=self.format_settings['bold'])
        self.italic_var = tk.BooleanVar(value=self.format_settings['italic'])
        self.underline_var = tk.BooleanVar(
            value=self.format_settings['underline'])

        ttk.Checkbutton(style_frame, text="Bold",
                        variable=self.bold_var).pack(anchor="w")
        ttk.Checkbutton(style_frame, text="Italic",
                        variable=self.italic_var).pack(anchor="w")
        ttk.Checkbutton(style_frame, text="Underline",
                        variable=self.underline_var).pack(anchor="w")

        # Preview Frame
        preview_frame = ttk.LabelFrame(self.dialog, text="Preview", padding=10)
        preview_frame.pack(fill="x", padx=20, pady=10)

        self.preview_label = ttk.Label(preview_frame, text="Sample Text")
        self.preview_label.pack(pady=10)

        # Buttons
        button_frame = ttk.Frame(self.dialog)
        button_frame.pack(fill="x", padx=20, pady=20)

        ttk.Button(button_frame, text="Apply",
                   command=self.apply_format).pack(side="right", padx=5)
        ttk.Button(button_frame, text="Cancel",
                   command=self.dialog.destroy).pack(side="right", padx=5)

        # Bind events for live preview
        self.font_combo.bind('<<ComboboxSelected>>',
                             lambda e: self.update_preview())
        self.size_var.trace_add('write', lambda *args: self.update_preview())
        self.bold_var.trace_add('write', lambda *args: self.update_preview())
        self.italic_var.trace_add('write', lambda *args: self.update_preview())
        self.underline_var.trace_add(
            'write', lambda *args: self.update_preview())

        self.update_preview()

    def choose_color(self):
        color = colorchooser.askcolor(self.format_settings['font_color'])[1]
        if color:
            self.format_settings['font_color'] = color
            self.update_color_button()
            self.update_preview()

    def update_color_button(self):
        self.color_button.configure(text=self.format_settings['font_color'])

    def update_preview(self):
        try:
            # Create font string
            font_style = []
            if self.bold_var.get():
                font_style.append('bold')
            if self.italic_var.get():
                font_style.append('italic')
            if self.underline_var.get():
                font_style.append('underline')

            # Update preview label
            self.preview_label.configure(
                font=(self.font_combo.get(), int(
                    self.size_var.get()), ' '.join(font_style)),
                foreground=self.format_settings['font_color']
            )
        except Exception:
            pass

    def apply_format(self):
        self.format_settings.update({
            'font_name': self.font_combo.get(),
            'font_size': int(self.size_var.get()),
            'bold': self.bold_var.get(),
            'italic': self.italic_var.get(),
            'underline': self.underline_var.get()
        })
        self.dialog.destroy()

    def get_format(self):
        return self.format_settings


class DocumentAutomation:
    def __init__(self):
        self.root = tk.Tk()
        self.root.title("Document Automation Tool")
        self.root.geometry("900x700")

        # Variables
        self.template_path = tk.StringVar()
        self.list_path = tk.StringVar()
        self.save_location = tk.StringVar()
        self.keywords = []
        self.template_keywords = []
        self.list_columns = []
        self.keyword_checkboxes = {}  # Store checkboxes for keywords
        self.keyword_formats = {}     # Store format settings for keywords

        # Output format checkboxes
        self.output_formats = {
            "pdf": tk.BooleanVar(value=True),
            "docx": tk.BooleanVar(value=True),
        }

        # Create frames for each step
        self.upload_frame = self.create_upload_frame()
        self.keyword_frame = self.create_keyword_frame()
        self.output_frame = self.create_output_frame()

        # Start with upload frame
        self.show_upload_frame()

    def detect_keywords(self, text):
        """
        Detect keywords from text with various formats.
        Rules:
        - Ignore email addresses
        - Only detect digits when wrapped in symbols
        - Brackets must have both opening and closing symbols
        """
        patterns = [
            # Double symbol with closing pair
            (r'\$\$(.+?)\$\$', '$$', '$$'),       # $$keyword$$ 
            (r'##(.+?)##', '##', '##'),           # ##keyword##
            (r'@@(.+?)@@', '@@', '@@'),           # @@keyword@@
            (r'\|\|(.+?)\|\|', '||', '||'),       # ||keyword||
            
            # Single symbol with closing pair
            (r'\$(.+?)\$', '$', '$'),             # $keyword$
            (r'#(.+?)#', '#', '#'),               # #keyword#
            (r'@(.+?)@', '@', '@'),               # @keyword@
            (r'\|(.+?)\|', '|', '|'),             # |keyword|
            
            # Double symbol without closing
            (r'\$\$(.+?)(?:\s|$)', '$$', None),   # $$keyword
            (r'\|\|(.+?)(?:\s|$)', '||', None),   # ||keyword
            (r'@@(.+?)(?:\s|$)', '@@', None),     # @@keyword
            
            # Single symbol without closing
            (r'\$(.+?)(?:\s|$)', '$', None),      # $keyword
            (r'\|(.+?)(?:\s|$)', '|', None),      # |keyword
            (r'@(.+?)(?:\s|$)', '@', None),       # @keyword
            
            # Double brackets
            (r'\{\{(.+?)\}\}', '{{', '}}'),       # {{keyword}}
            (r'\[\[(.+?)\]\]', '[[', ']]'),       # [[keyword]]
            (r'\(\((.+?)\)\)', '((', '))'),       # ((keyword))
            
            # Single brackets
            (r'\{(.+?)\}', '{', '}'),             # {keyword}
            (r'\[(.+?)\]', '[', ']'),             # [keyword]
            (r'\((.+?)\)', '(', ')'),             # (keyword)
            
            # Mixed formats - brackets with symbols
            (r'\{\$(.+?)\$\}', '{$', '$}'),       # {$keyword$}
            (r'\{#(.+?)#\}', '{#', '#}'),         # {#keyword#}
            (r'\[#(.+?)#\]', '[#', '#]'),         # [#keyword#]
            (r'\[\$(.+?)\$\]', '[$', '$]'),       # [$keyword$]
            (r'\(#(.+?)#\)', '(#', '#)'),         # (#keyword#)
            (r'\(\$(.+?)\$\)', '($', '$)'),       # ($keyword$)
            
            # Double symbols with brackets
            (r'\{\$\$(.+?)\$\$\}', '{$$', '$$}'), # {$$keyword$$}
            (r'\[##(.+?)##\]', '[##', '##]'),     # [##keyword##]
            (r'\(##(.+?)##\)', '(##', '##)'),     # (##keyword##)
        ]

        # Email pattern for filtering
        email_pattern = r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}'
        email_addresses = set(re.findall(email_pattern, text))

        keywords = []
        self.keyword_symbols = {}  # Store original symbols for each keyword

        def is_valid_keyword(keyword, raw_match):
            """
            Validate if keyword meets the criteria:
            - Not part of an email address
            - Contains only digits when wrapped in symbols
            """
            # Remove extra whitespace
            keyword = keyword.strip()
            if not keyword:
                return False
                
            # Check if the keyword is part of an email address
            if any(keyword in email for email in email_addresses):
                return False
                
            # Check if the raw match (including symbols) is an email address
            if raw_match in email_addresses:
                return False
                
            # Only accept digits
            return keyword.isdigit()

        for pattern, start_symbol, end_symbol in patterns:
            matches = re.finditer(pattern, text)
            for match in matches:
                # Get the keyword inside the symbols
                keyword = match.group(1).strip()
                raw_match = match.group(0).strip()
                
                # Only add keyword if it meets the validation criteria
                if is_valid_keyword(keyword, raw_match):
                    if keyword not in self.keyword_symbols:
                        self.keyword_symbols[keyword] = []
                    # Only store if it's a complete bracket pair or a special symbol
                    if (start_symbol and end_symbol) or start_symbol in ['$$', '||', '@@', '$', '|', '@']:
                        self.keyword_symbols[keyword].append((start_symbol, end_symbol))

        return list(set(self.keyword_symbols.keys()))  # Return unique keywords

    def create_upload_frame(self):
        """Create the first frame for file upload"""
        frame = ttk.Frame(self.root)

        title = ttk.Label(frame, text="Step 1: Upload Files",
                          font=("Helvetica", 14, "bold"))
        title.pack(pady=20)

        # Info label
        info_text = "Please upload your template file (.docx) and list file (.xlsx or .csv) with keywords with symbols or brackets"
        ttk.Label(frame, text=info_text, wraplength=600).pack(pady=10)

        # Template file
        template_group = ttk.LabelFrame(
            frame, text="Template File (.docx)", padding=10)
        template_group.pack(fill="x", padx=20, pady=10)

        ttk.Entry(template_group, textvariable=self.template_path,
                  width=60).pack(side="left", padx=5)
        ttk.Button(template_group, text="Browse",
                   command=self.browse_template).pack(side="left")

        # Template keywords preview
        self.template_preview = ttk.Label(frame, text="", wraplength=600)
        self.template_preview.pack(pady=10)

        # List file
        list_group = ttk.LabelFrame(
            frame, text="List File (.xlsx, .csv)", padding=10)
        list_group.pack(fill="x", padx=20, pady=10)

        ttk.Entry(list_group, textvariable=self.list_path,
                  width=60).pack(side="left", padx=5)
        ttk.Button(list_group, text="Browse",
                   command=self.browse_list).pack(side="left")

        # List columns preview
        self.list_preview = ttk.Label(frame, text="", wraplength=600)
        self.list_preview.pack(pady=10)

        # Next button
        ttk.Button(frame, text="Check Keywords and Continue →",
                   command=self.check_and_proceed).pack(pady=20)

        return frame

    def create_keyword_frame(self):
        """Create the second frame for keyword matching"""
        frame = ttk.Frame(self.root)

        title = ttk.Label(frame, text="Step 2: Keyword Matching Results",
                        font=("Helvetica", 14, "bold"))
        title.pack(pady=20)

        # Results section with scrollbar
        results_container = ttk.Frame(frame)
        results_container.pack(fill="both", expand=True, padx=20, pady=10)

        # Add scrollbar
        scrollbar = ttk.Scrollbar(results_container)
        scrollbar.pack(side="right", fill="y")

        # Create canvas for scrolling
        canvas = tk.Canvas(results_container, yscrollcommand=scrollbar.set)
        canvas.pack(side="left", fill="both", expand=True)

        scrollbar.config(command=canvas.yview)

        # Create frame for results inside canvas
        self.results_frame = ttk.Frame(canvas)
        canvas.create_window((0, 0), window=self.results_frame, anchor="nw")

        # Configure scroll region when results frame changes
        self.results_frame.bind("<Configure>",
                                lambda e: canvas.configure(scrollregion=canvas.bbox("all")))

        # Summary frame
        self.summary_frame = ttk.LabelFrame(
            frame, text="Matching Summary", padding=10)
        self.summary_frame.pack(fill="x", padx=20, pady=10)

        # Navigation buttons
        nav_frame = ttk.Frame(frame)
        nav_frame.pack(fill="x", padx=20, pady=20)
        ttk.Button(nav_frame, text="← Back",
                command=self.show_upload_frame).pack(side="left")
        self.continue_button = ttk.Button(nav_frame,
                                        text="Continue to Output Settings →",
                                        command=self.show_output_frame,
                                        state="disabled")
        self.continue_button.pack(side="right")

        return frame

    def toggle_all_keywords(self, select_all):
        """Toggle all keyword checkboxes"""
        for var in self.keyword_checkboxes.values():
            var.set(select_all)
        self.check_selected_keywords()

    def auto_match_keywords(self, list_columns, template_keywords):
        """Enhanced automatic matching of keywords with columns, especially for numeric patterns"""
        matches = {}
        
        for keyword in template_keywords:
            # Extract numbers from keyword
            keyword_numbers = ''.join(filter(str.isdigit, keyword))
            
            # Try different matching strategies
            for column in list_columns:
                column_numbers = ''.join(filter(str.isdigit, column))
                
                # Strategy 1: Exact number match
                if keyword_numbers and column_numbers and keyword_numbers == column_numbers:
                    matches[keyword] = column
                    break
                    
                # Strategy 2: Exact string match (case insensitive)
                elif column.lower() == keyword.lower():
                    matches[keyword] = column
                    break
                    
                # Strategy 3: Number at the end match
                elif keyword_numbers and column_numbers and keyword_numbers == column_numbers[-len(keyword_numbers):]:
                    matches[keyword] = column
                    break
                    
                # Strategy 4: Partial number match
                elif keyword_numbers and column_numbers and (keyword_numbers in column_numbers or column_numbers in keyword_numbers):
                    matches[keyword] = column
                    break

        return matches

    def clean_string(self, text):
        """Clean string for better matching"""
        # Convert to lowercase
        text = text.lower()
        # Remove common symbols and separators
        text = text.replace('_', '').replace('-', '').replace(' ', '')
        # Remove common prefixes/suffixes
        prefixes = ['$$', '##', '{{', '}}', '[[', ']]', '(', ')', '{', '}', '|']
        for prefix in prefixes:
            text = text.replace(prefix, '')
        return text

    def create_output_frame(self):
        """Create the third frame for output settings"""
        frame = ttk.Frame(self.root)

        title = ttk.Label(frame, text="Step 3: Output Settings",
                          font=("Helvetica", 14, "bold"))
        title.pack(pady=20)

        # Output format checkboxes
        format_group = ttk.LabelFrame(
            frame, text="Select Output Format(s)", padding=10)
        format_group.pack(fill="x", padx=20, pady=10)

        ttk.Checkbutton(format_group, text="PDF Document",
                        variable=self.output_formats["pdf"]).pack(anchor="w")
        ttk.Checkbutton(format_group, text="Word Document",
                        variable=self.output_formats["docx"]).pack(anchor="w")

        # Save location
        location_group = ttk.LabelFrame(
            frame, text="Save Location", padding=10)
        location_group.pack(fill="x", padx=20, pady=10)

        ttk.Entry(location_group, textvariable=self.save_location,
                  width=60).pack(side="left", padx=5)
        ttk.Button(location_group, text="Browse",
                   command=self.browse_save_location).pack(side="left")

        # Navigation buttons
        nav_frame = ttk.Frame(frame)
        nav_frame.pack(fill="x", padx=20, pady=20)
        ttk.Button(nav_frame, text="← Back",
                   command=self.show_keyword_frame).pack(side="left")
        ttk.Button(nav_frame, text="Generate Files",
                   command=self.process_files).pack(side="right")

        return frame

    def edit_keyword_format(self, keyword):
        """Open format dialog for keyword"""
        current_format = self.keyword_formats.get(keyword, None)
        dialog = KeywordFormatDialog(self.root, current_format)
        self.root.wait_window(dialog.dialog)
        new_format = dialog.get_format()
        if new_format:
            self.keyword_formats[keyword] = new_format
            self.update_format_preview(keyword)

    def update_format_preview(self, keyword):
        """Update the format preview for a keyword"""
        if keyword in self.keyword_formats:
            format_settings = self.keyword_formats[keyword]
            preview_text = f"{format_settings['font_name']}, {
                format_settings['font_size']}pt"
            if format_settings['bold']:
                preview_text += ", B"
            if format_settings['italic']:
                preview_text += ", I"
            if format_settings['underline']:
                preview_text += ", U"
            return preview_text
        return "Default"

    def find_name_column(self, columns):
        """Find the name column using case-insensitive comparison"""
        # Check for various possible name column formats
        name_variants = ['name', 'names', 'full name',
                         'fullname', '$$name', 'NAME', '$$NAME']
        for column in columns:
            # Convert to lowercase for comparison
            col_lower = column.lower()
            # Check if the column contains any of the name variants
            if any(variant.lower() in col_lower for variant in name_variants):
                return column
        return None

    def find_folder_column(self, columns):
        """Find the folder column regardless of case"""
        folder_variants = ['folder', 'folder_name', 'foldername']
        for column in columns:
            if column.lower() in folder_variants:
                return column
        return None

    def auto_match_keywords(self, list_columns, template_keywords):
        """Automatically match keywords with columns using case-insensitive comparison"""
        matches = {}
        for keyword in template_keywords:
            keyword_lower = keyword.lower()
            # Try exact match first
            if keyword in list_columns:
                matches[keyword] = keyword
            else:
                # Try case-insensitive match
                for column in list_columns:
                    if column.lower() == keyword_lower:
                        matches[keyword] = column
                        break
        return matches

    def check_and_proceed(self):
        """Validate files and proceed to keyword matching step"""
        if not self.template_path.get() or not self.list_path.get():
            messagebox.showerror("Error", "Please select both template and list files")
            return

        # Clear previous results
        for widget in self.results_frame.winfo_children():
            widget.destroy()
        for widget in self.summary_frame.winfo_children():
            widget.destroy()

        # Get keywords and columns
        template_keywords = self.detect_template_keywords()
        list_columns = self.detect_list_columns()

        # Auto-match keywords
        auto_matches = self.auto_match_keywords(list_columns, template_keywords)

        # Create header container frames
        headers = ["Select All", "Keyword", "Format(s)", "Status", "Match", "Format"]
        header_frames = []
        
        for col, header in enumerate(headers):
            # Create a frame for each header to contain both checkbox and label if needed
            header_frame = ttk.Frame(self.results_frame)
            header_frame.grid(row=0, column=col, padx=5, pady=5)
            header_frames.append(header_frame)
            
            if col == 0:  # Special handling for "Select All" header
                # Add checkbox
                select_all_var = tk.BooleanVar(value=False)
                select_all_chk = ttk.Checkbutton(
                    header_frame,
                    variable=select_all_var,
                    command=lambda: self.toggle_all_keywords(select_all_var.get())
                )
                select_all_chk.pack(side="left")
                
                # Add label
                ttk.Label(
                    header_frame,
                    text="Select All",
                    font=("Helvetica", 10, "bold")
                ).pack(side="left")
            else:
                # Regular headers
                ttk.Label(
                    header_frame,
                    text=header,
                    font=("Helvetica", 10, "bold")
                ).pack()

        self.keyword_checkboxes.clear()
        self.keywords = []
        self.status_labels = {}  # Store status labels for updating

        def is_exact_match(keyword, column):
            """Check if keyword exactly matches the column name"""
            # Remove symbols and compare
            keyword_clean = self.clean_string(keyword)
            column_clean = self.clean_string(column)
            return keyword_clean == column_clean

        for i, keyword in enumerate(template_keywords, 1):
            # Check for exact match in auto_matches
            is_matched = False
            if keyword in auto_matches:
                is_matched = is_exact_match(keyword, auto_matches[keyword])

            # Checkbox (only pre-selected if exact match)
            var = tk.BooleanVar(value=is_matched)
            self.keyword_checkboxes[keyword] = var
            chk = ttk.Checkbutton(self.results_frame, variable=var, command=self.check_selected_keywords)
            chk.grid(row=i, column=0, padx=5, pady=2)

            # Keyword
            ttk.Label(self.results_frame, text=keyword).grid(row=i, column=1, padx=5, pady=2)

            # Format column
            formats = []
            for start_symbol, end_symbol in self.keyword_symbols[keyword]:
                if end_symbol:
                    formats.append(f"{start_symbol}...{end_symbol}")
                else:
                    formats.append(f"{start_symbol}")
            ttk.Label(self.results_frame, text=", ".join(formats)).grid(row=i, column=2, padx=5, pady=2)

            # Status and Match
            status_label = ttk.Label(self.results_frame, text="× Not matched", foreground="orange")
            status_label.grid(row=i, column=3, padx=5, pady=2)
            self.status_labels[keyword] = status_label

            combo = ttk.Combobox(self.results_frame, values=list_columns, state="readonly", width=30)
            combo.grid(row=i, column=4, padx=5, pady=2)
            
            # Set auto-matched value if exists
            if keyword in auto_matches:
                combo.set(auto_matches[keyword])
                self.update_status(keyword, auto_matches[keyword])

            # Bind the combobox selection event
            combo.bind('<<ComboboxSelected>>', lambda e, k=keyword, c=combo: self.check_match(k, c))

            # Format button
            format_button = ttk.Button(
                self.results_frame,
                text="Edit Format",
                command=lambda k=keyword: self.edit_keyword_format(k)
            )
            format_button.grid(row=i, column=5, padx=5, pady=2)

            # Format preview label
            format_preview = ttk.Label(self.results_frame, text=self.update_format_preview(keyword))
            format_preview.grid(row=i, column=6, padx=5, pady=2)

            self.keywords.append((keyword, combo))

        self.check_selected_keywords()
        self.show_keyword_frame()

    def check_match(self, keyword, combo):
        """Check if selected column matches the keyword and update status and checkbox"""
        selected_column = combo.get()
        if selected_column:
            keyword_clean = self.clean_string(keyword)
            column_clean = self.clean_string(selected_column)
            
            # Check for matches
            is_match = False
            match_type = ""

            # Exact match (case-insensitive)
            if keyword_clean == column_clean:
                is_match = True
                match_type = "Exact match"
                # Auto-check the checkbox for exact matches
                self.keyword_checkboxes[keyword].set(True)
            
            # Number match
            elif self.extract_numbers(keyword_clean) == self.extract_numbers(column_clean):
                is_match = True
                match_type = "Number match"
            
            # Partial text match
            elif keyword_clean in column_clean or column_clean in keyword_clean:
                is_match = True
                match_type = "Text match"

            # Update status label
            if is_match:
                self.status_labels[keyword].config(
                    text=f"✓ {match_type}",
                    foreground="green"
                )
            else:
                self.status_labels[keyword].config(
                    text="× Not matched",
                    foreground="orange"
                )
                
            # After status update, check if we need to update the "Select All" state
            self.check_selected_keywords()
    
    def extract_numbers(self, text):
        """Extract only numbers from text"""
        return ''.join(filter(str.isdigit, text))

    def update_status(self, keyword, column):
        """Update the status label based on match detection"""
        keyword_clean = self.clean_string(keyword)
        column_clean = self.clean_string(column)

        # Check for matches
        is_match = False
        match_type = ""

        # Exact match (case-insensitive)
        if keyword_clean == column_clean:
            is_match = True
            match_type = "Exact match"
        
        # Number match
        elif self.extract_numbers(keyword_clean) == self.extract_numbers(column_clean):
            is_match = True
            match_type = "Number match"
        
        # Partial text match
        elif keyword_clean in column_clean or column_clean in keyword_clean:
            is_match = True
            match_type = "Text match"

        # Update status label
        if is_match:
            self.status_labels[keyword].config(
                text=f"✓ {match_type}",
                foreground="green"
            )
        else:
            self.status_labels[keyword].config(
                text="× Not matched",
                foreground="orange"
            )

    def clean_string(self, text):
        """Clean string for matching by removing symbols and converting to lowercase"""
        # Convert to lowercase
        text = text.lower().strip()
        # Remove common symbols and separators
        text = text.replace('_', '').replace('-', '').replace(' ', '')
        # Remove common prefixes/suffixes
        prefixes = ['$$', '##', '{{', '}}', '[[', ']]', '(', ')', '{', '}', '|']
        for prefix in prefixes:
            text = text.replace(prefix, '')
        return text

    def check_selected_keywords(self):
        """Update summary and continue button based on selected keywords and their matches"""
        # Count selected keywords
        selected_count = 0
        unmatched_keywords = []
        
        # Check each selected keyword for a match
        for keyword, var in self.keyword_checkboxes.items():
            if var.get():  # If checkbox is selected
                selected_count += 1
                # Find corresponding combobox
                matched = False
                for kw, combo in self.keywords:
                    if kw == keyword:
                        if not combo.get():  # If no match selected
                            unmatched_keywords.append(keyword)
                        break

        total_keywords = len(self.keyword_checkboxes)

        # Update summary
        for widget in self.summary_frame.winfo_children():
            widget.destroy()

        summary_text = f"Selected {selected_count} out of {total_keywords} keywords"
        ttk.Label(self.summary_frame, text=summary_text).pack()

        # Enable/disable continue button based on selection and matches
        if selected_count > 0 and not unmatched_keywords:
            self.continue_button.config(state="normal")
            message = "You can proceed to output settings."
            ttk.Label(self.summary_frame, text=message, foreground="green").pack()
        else:
            self.continue_button.config(state="disabled")
            if selected_count == 0:
                message = "Please select at least one keyword to proceed."
            else:
                message = f"Please select matches for: {', '.join(unmatched_keywords)}"
            ttk.Label(self.summary_frame, text=message, foreground="orange").pack()

    def detect_template_keywords(self):
        """Detect keywords from template document, including all possible text locations"""
        if not self.template_path.get():
            return []

        try:
            doc = Document(self.template_path.get())
            all_text = []

            # Process regular paragraphs
            for paragraph in doc.paragraphs:
                all_text.append(paragraph.text)

            # Process tables and nested tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        # Handle paragraphs within cells
                        for paragraph in cell.paragraphs:
                            if paragraph.text.strip():
                                all_text.append(paragraph.text)

                        # Handle nested tables
                        for nested_table in cell.tables:
                            for nested_row in nested_table.rows:
                                for nested_cell in nested_row.cells:
                                    for nested_para in nested_cell.paragraphs:
                                        if nested_para.text.strip():
                                            all_text.append(nested_para.text)

            # Process shapes and text boxes
            for shape in doc.inline_shapes:
                if hasattr(shape, 'text_frame'):
                    for paragraph in shape.text_frame.paragraphs:
                        if paragraph.text.strip():
                            all_text.append(paragraph.text)

            # Process floating shapes and text boxes (in document._element.body)
            for shape in doc._element.findall('.//w:txbxContent/w:p', doc._element.nsmap):
                text = shape.xpath('string()')
                if text.strip():
                    all_text.append(text)

            # Combine all text and detect keywords
            text = '\n'.join(all_text)
            keywords = self.detect_keywords(text)

            # Update preview
            if keywords:
                preview_text = "Found keywords: " + ", ".join(keywords)
            else:
                preview_text = "No keywords found in template"
            self.template_preview.config(text=preview_text)

            return keywords

        except Exception as e:
            messagebox.showerror(
                "Error", f"Error reading template file:\n{str(e)}")
            return []

    def detect_list_columns(self):
        """Detect column names from list file"""
        if not self.list_path.get():
            return []

        try:
            # Read file
            if self.list_path.get().endswith('.xlsx'):
                df = pd.read_excel(self.list_path.get())
            else:
                df = pd.read_csv(self.list_path.get())

            columns = list(df.columns)

            # Update preview
            if columns:
                preview_text = "Found columns: " + ", ".join(columns)
            else:
                preview_text = "No columns found in list file"
            self.list_preview.config(text=preview_text)

            return columns

        except Exception as e:
            messagebox.showerror(
                "Error", f"Error reading list file:\n{str(e)}")
            return []

    def process_files(self):
        """Process the files with selected keywords and output formats"""
        if not self.save_location.get():
            messagebox.showerror("Error", "Please select a save location")
            return

        if not any(format_var.get() for format_var in self.output_formats.values()):
            messagebox.showerror(
                "Error", "Please select at least one output format")
            return

        try:
            # Read data
            df = pd.read_excel(self.list_path.get()) if self.list_path.get().endswith(
                '.xlsx') else pd.read_csv(self.list_path.get())

            # Find name and folder columns
            name_column = self.find_name_column(df.columns)
            folder_column = self.find_folder_column(df.columns)

            if not name_column:
                messagebox.showerror(
                    "Error", "Name column not found in list file. Please ensure you have a column with 'name' in it (case insensitive).")
                return

            # Create mapping from keywords to column names
            mapping = {}
            for keyword, var in self.keyword_checkboxes.items():
                if var.get():  # Only process selected keywords
                    for kw, combo in self.keywords:
                        if kw == keyword and combo.get():
                            mapping[keyword] = combo.get()

            # Create progress window
            progress_window = tk.Toplevel(self.root)
            progress_window.title("Processing Files")
            progress_window.geometry("400x150")
            progress_window.transient(self.root)

            # Center progress window
            progress_window.geometry("+%d+%d" % (
                self.root.winfo_x() + self.root.winfo_width()/2 - 200,
                self.root.winfo_y() + self.root.winfo_height()/2 - 75))

            # Add progress label and bar
            progress_label = ttk.Label(progress_window, text="Processing files...")
            progress_label.pack(pady=10)
            progress_bar = ttk.Progressbar(progress_window, length=300, mode='determinate')
            progress_bar.pack(pady=10)
            file_label = ttk.Label(progress_window, text="")
            file_label.pack(pady=10)

            # Process each row
            total_files = len(df)
            successful_files = 0
            failed_files = []
            progress_bar['maximum'] = total_files

            for index, row in df.iterrows():
                try:
                    # Update progress
                    progress_bar['value'] = index + 1
                    file_label.config(text=f"Processing file {index + 1} of {total_files}")
                    progress_window.update()

                    # Create new document from template
                    doc = Document(self.template_path.get())

                    # Process regular paragraphs
                    for paragraph in doc.paragraphs:
                        self._replace_keywords_in_paragraph(paragraph, row, mapping)

                    # Process tables and nested tables
                    for table in doc.tables:
                        for table_row in table.rows:
                            for cell in table_row.cells:
                                # Process paragraphs within cells
                                for paragraph in cell.paragraphs:
                                    self._replace_keywords_in_paragraph(paragraph, row, mapping)

                                # Process nested tables
                                for nested_table in cell.tables:
                                    for nested_row in nested_table.rows:
                                        for nested_cell in nested_row.cells:
                                            for nested_para in nested_cell.paragraphs:
                                                self._replace_keywords_in_paragraph(nested_para, row, mapping)

                    # Process shapes and text boxes
                    for shape in doc.inline_shapes:
                        if hasattr(shape, 'text_frame'):
                            for paragraph in shape.text_frame.paragraphs:
                                self._replace_keywords_in_paragraph(paragraph, row, mapping)

                    # Process floating shapes and text boxes
                    try:
                        for shape in doc._element.findall('.//w:txbxContent/w:p', doc._element.nsmap):
                            text = shape.xpath('string()')
                            if text.strip():
                                # Create a temporary paragraph object for processing
                                temp_paragraph = doc.add_paragraph()
                                temp_paragraph.text = text
                                self._replace_keywords_in_paragraph(temp_paragraph, row, mapping)
                                # Update the shape text with the processed text
                                shape.text = temp_paragraph.text
                                # Remove the temporary paragraph
                                temp_paragraph._element.getparent().remove(temp_paragraph._element)
                    except Exception as shape_error:
                        print(f"Warning: Error processing shapes in document: {str(shape_error)}")

                    # Get output name from name column - handle empty or invalid values
                    output_name = str(row[name_column]).strip()
                    if not output_name:
                        output_name = f"document_{index + 1}"
                    # Clean the output name
                    output_name = "".join(
                        x for x in output_name if x.isalnum() or x in (' ', '-', '_'))
                    if not output_name:  # If after cleaning the name is empty
                        output_name = f"document_{index + 1}"

                    # Get custom folder name or use default
                    folder_name = str(row[folder_column]) if folder_column and pd.notna(
                        row[folder_column]) else "default"
                    folder_name = "".join(
                        x for x in folder_name if x.isalnum() or x in (' ', '-', '_'))
                    if not folder_name:  # If after cleaning the folder name is empty
                        folder_name = "default"

                    # Create base directory for this record
                    record_base_dir = Path(self.save_location.get()) / folder_name

                    # Create format-specific subdirectories
                    output_dirs = {}
                    for format_type, format_var in self.output_formats.items():
                        if format_var.get():
                            output_dir = record_base_dir / format_type
                            output_dir.mkdir(parents=True, exist_ok=True)
                            output_dirs[format_type] = output_dir

                    # Save in selected formats
                    if self.output_formats["docx"].get():
                        docx_path = output_dirs["docx"] / f"{output_name}.docx"
                        doc.save(docx_path)

                    if self.output_formats["pdf"].get():
                        pdf_dir = output_dirs["pdf"]
                        temp_docx = pdf_dir / f"{output_name}.docx"
                        pdf_path = pdf_dir / f"{output_name}.pdf"

                        # Save temporary docx for PDF conversion
                        doc.save(temp_docx)
                        try:
                            convert(str(temp_docx), str(pdf_path))
                            # Remove temporary docx file after successful PDF conversion
                            temp_docx.unlink()
                        except Exception as pdf_error:
                            failed_files.append(
                                (output_name, f"PDF conversion error: {str(pdf_error)}"))
                            continue

                    successful_files += 1

                except Exception as row_error:
                    failed_files.append((output_name, str(row_error)))
                    continue

                # Update progress window
                progress_window.update()

            # Close progress window
            progress_window.destroy()

            # Show completion message
            completion_message = f"Processing complete!\n\n"
            completion_message += f"Successfully processed: {successful_files} files\n"

            if failed_files:
                completion_message += f"\nFailed to process {len(failed_files)} files:\n"
                # Show first 5 failures
                for failed_file, error in failed_files[:5]:
                    completion_message += f"- {failed_file}: {error}\n"
                if len(failed_files) > 5:
                    completion_message += f"(and {len(failed_files) - 5} more...)\n"

            completion_message += f"\nFiles have been saved to:\n{self.save_location.get()}" \
                                f"\n\nWould you like to open the output folder?"

            result = messagebox.askquestion(
                "Processing Complete",
                completion_message,
                icon='info'
            )

            # Open folder if requested
            if result == 'yes':
                os.startfile(str(self.save_location.get()))

            # Reset the form
            self.template_path.set("")
            self.list_path.set("")
            self.save_location.set("")
            self.template_preview.config(text="")
            self.list_preview.config(text="")
            self.keyword_formats.clear()

            # Return to first page
            self.show_upload_frame()

        except Exception as e:
            messagebox.showerror(
                "Error",
                f"An error occurred during processing:\n{str(e)}\n\n"
                "Please check your input files and try again."
            )
            return

    def _replace_keywords_in_paragraph(self, paragraph, row, mapping):
        """Helper method to replace keywords in a paragraph with proper formatting"""
        for keyword in mapping:
            if keyword in self.keyword_symbols:
                for start_symbol, end_symbol in self.keyword_symbols[keyword]:
                    if end_symbol:
                        original = start_symbol + keyword + end_symbol
                    else:
                        original = start_symbol + keyword

                    if original in paragraph.text:
                        new_value = str(row[mapping[keyword]])

                        # Find the run containing the keyword and apply formatting
                        for run in paragraph.runs:
                            if original in run.text:
                                # Apply formatting if specified
                                if keyword in self.keyword_formats:
                                    format_settings = self.keyword_formats[keyword]
                                    run.font.name = format_settings['font_name']
                                    run.font.size = Pt(
                                        format_settings['font_size'])
                                    run.font.color.rgb = RGBColor.from_string(
                                        format_settings['font_color'][1:])
                                    run.font.bold = format_settings['bold']
                                    run.font.italic = format_settings['italic']
                                    run.font.underline = format_settings['underline']

                                # Replace text
                                run.text = run.text.replace(
                                    original, new_value)

    def browse_template(self):
        """Open file dialog for template selection"""
        filename = filedialog.askopenfilename(
            filetypes=[("Word files", "*.docx"), ("All files", "*.*")]
        )
        if filename:
            self.template_path.set(filename)
            self.detect_template_keywords()

    def browse_list(self):
        """Open file dialog for list file selection"""
        filename = filedialog.askopenfilename(
            filetypes=[
                ("Excel files", "*.xlsx"),
                ("CSV files", "*.csv"),
                ("All files", "*.*")
            ]
        )
        if filename:
            self.list_path.set(filename)
            self.detect_list_columns()

    def browse_save_location(self):
        """Open directory dialog for save location selection"""
        directory = filedialog.askdirectory()
        if directory:
            self.save_location.set(directory)

    def show_upload_frame(self):
        """Switch to upload frame"""
        self.keyword_frame.pack_forget()
        self.output_frame.pack_forget()
        self.upload_frame.pack(fill="both", expand=True)

    def show_keyword_frame(self):
        """Switch to keyword frame"""
        self.upload_frame.pack_forget()
        self.output_frame.pack_forget()
        self.keyword_frame.pack(fill="both", expand=True)

    def show_output_frame(self):
        """Switch to output frame"""
        if not any(var.get() for var in self.keyword_checkboxes.values()):
            messagebox.showerror(
                "Error", "Please select at least one keyword to proceed")
            return

        self.upload_frame.pack_forget()
        self.keyword_frame.pack_forget()
        self.output_frame.pack(fill="both", expand=True)

    def run(self):
        """Start the application"""
        self.root.mainloop()


if __name__ == "__main__":
    try:
        app = DocumentAutomation()
        app.run()
    except Exception as e:
        messagebox.showerror("Critical Error",
                             f"An unexpected error occurred while starting the application:\n{
                                 str(e)}\n\n"
                             "Please ensure all required dependencies are installed:\n"
                             "- python-docx\n"
                             "- docx2pdf\n"
                             "- pandas\n"
                             "- openpyxl"
                             )
